# 文件上传适配器
from typing import List, Dict, Any
from langchain.schema import Document
from pathlib import Path
import tempfile
import logging
import io

from .base import DocumentAdapter
from app.services.storage.minio_service import get_minio_service

logger = logging.getLogger(__name__)


class FileUploadAdapter(DocumentAdapter):
    """MinIO 上传文件适配器

    支持从 MinIO 加载并解析文件：
    - PDF: 使用 pypdf 解析
    - TXT/MD: 直接读取文本
    - DOCX: 使用 python-docx 解析
    """

    def __init__(self, file_paths: List[str]):
        """初始化

        Args:
            file_paths: MinIO 文件路径列表（object_name）
        """
        self.file_paths = file_paths
        self.minio_service = get_minio_service()

    async def load(self) -> List[Document]:
        """从 MinIO 加载文件并解析为文档"""
        documents = []

        for file_path in self.file_paths:
            try:
                # 从 MinIO 下载文件
                download_result = await self.minio_service.download_file(
                    bucket="documents",
                    object_name=file_path
                )

                if not download_result.get("success"):
                    logger.error(f"下载文件失败: {file_path}, {download_result.get('error')}")
                    continue

                file_data = download_result["data"]
                file_ext = Path(file_path).suffix.lower()

                # 根据文件类型解析
                parsed_docs = await self._parse_file(file_data, file_ext, file_path)
                documents.extend(parsed_docs)

            except Exception as e:
                logger.error(f"处理文件 {file_path} 失败: {e}")
                continue

        return documents

    async def _parse_file(
        self,
        file_data: bytes,
        file_ext: str,
        file_path: str
    ) -> List[Document]:
        """根据文件类型解析内容

        Args:
            file_data: 文件二进制数据
            file_ext: 文件扩展名
            file_path: 文件路径（用于元数据）

        Returns:
            解析后的文档列表
        """
        documents = []
        metadata = {"source": file_path, "file_type": file_ext}

        try:
            if file_ext == ".pdf":
                # PDF 解析
                pages = self._parse_pdf(file_data)
                for i, page_text in enumerate(pages):
                    if page_text.strip():
                        doc = Document(
                            page_content=page_text,
                            metadata={**metadata, "page": i + 1}
                        )
                        documents.append(doc)

            elif file_ext in [".txt", ".md"]:
                # 文本文件直接读取
                text = file_data.decode("utf-8", errors="ignore")
                if text.strip():
                    doc = Document(page_content=text, metadata=metadata)
                    documents.append(doc)

            elif file_ext == ".docx":
                # DOCX 解析
                paragraphs = self._parse_docx(file_data)
                text = "\n\n".join(paragraphs)
                if text.strip():
                    doc = Document(page_content=text, metadata=metadata)
                    documents.append(doc)

            else:
                logger.warning(f"不支持的文件类型: {file_ext}")

        except Exception as e:
            logger.error(f"解析文件 {file_path} 失败: {e}")

        return documents

    def _parse_pdf(self, file_data: bytes) -> List[str]:
        """解析 PDF 文件"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_data))
            pages = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            return pages

        except ImportError:
            logger.warning("pypdf 未安装，无法解析 PDF")
            return []

    def _parse_docx(self, file_data: bytes) -> List[str]:
        """解析 DOCX 文件"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_data))
            paragraphs = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            return paragraphs

        except ImportError:
            logger.warning("python-docx 未安装，无法解析 DOCX")
            return []

    def get_source_info(self) -> Dict[str, Any]:
        """获取源信息"""
        return {
            "type": "file_upload",
            "file_count": len(self.file_paths),
            "file_paths": self.file_paths
        }